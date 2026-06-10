"""Chuyển đổi text/markdown từ draft_document thành file .docx chuẩn pháp lý VN.

Định dạng áp dụng:
  - Font: Times New Roman 13pt (chuẩn văn bản nhà nước VN)
  - Lề: trái 3cm, phải 2cm, trên/dưới 2.5cm
  - Quốc hiệu / tiêu đề → căn giữa
  - **bold** → in đậm
  - [THÔNG TIN CẦN ĐIỀN] → màu cam, in đậm (dễ nhận ra chỗ cần điền)
  - ─── SECTION ─── → header section in đậm màu xanh
  - - item / • item → danh sách gạch đầu dòng
  - | bảng | → bảng word
"""
from __future__ import annotations

import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── Regex patterns ─────────────────────────────────────────────────────────────

_BOLD_RE       = re.compile(r"\*\*(.+?)\*\*")
_PLACEHOLDER_RE = re.compile(r"(\[[^\]]{3,60}\])")   # [THÔNG TIN CẦN ĐIỀN]
_SECTION_RE    = re.compile(r"^─{2,}\s*(.+?)\s*─{2,}$")
_SEPARATOR_RE  = re.compile(r"^[─\-_=]{3,}\s*$")
_H2_RE         = re.compile(r"^#{1,2}\s+(.+)$")
_H3_RE         = re.compile(r"^###\s+(.+)$")
_BULLET_RE     = re.compile(r"^[\-•]\s+(.+)$")
_TABLE_ROW_RE  = re.compile(r"^\|.+\|$")
_TABLE_SEP_RE  = re.compile(r"^\|[\s\-:|]+\|$")

# Từ khoá nhận diện dòng căn giữa (quốc hiệu, tiêu đề văn bản)
_CENTER_TRIGGERS = [
    "CỘNG HÒA XÃ HỘI CHỦ NGHĨA", "ĐỘC LẬP – TỰ DO", "ĐỘC LẬP - TỰ DO",
    "ĐỘC LẬP – TỰ DO – HẠNH PHÚC",
]
_TITLE_KEYWORDS = [
    "BIÊN BẢN", "HỢP ĐỒNG", "QUYẾT ĐỊNH", "CÔNG VĂN", "THÔNG BÁO",
    "ĐƠN XIN", "ĐƠN KHIẾU NẠI", "ĐƠN TỐ CÁO", "DI CHÚC", "GIẤY",
]

# Màu sắc
_COLOR_BLUE      = RGBColor(0x1F, 0x4E, 0x79)   # tiêu đề section
_COLOR_ORANGE    = RGBColor(0xC0, 0x40, 0x00)   # placeholder [...]
_COLOR_GRAY      = RGBColor(0x90, 0x90, 0x90)   # dòng phân cách


# ── Helpers ────────────────────────────────────────────────────────────────────

def _is_center_line(line: str) -> bool:
    up = line.upper()
    return any(t in up for t in _CENTER_TRIGGERS)


def _is_title_line(line: str) -> bool:
    stripped = line.strip()
    if stripped != stripped.upper() or len(stripped) < 5:
        return False
    if _SEPARATOR_RE.match(stripped) or stripped.startswith("|"):
        return False
    return any(kw in stripped for kw in _TITLE_KEYWORDS)


def _add_formatted_runs(para, text: str) -> None:
    """Thêm runs vào paragraph, xử lý **bold** và [placeholder]."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\[[^\]]{3,60}\])")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("[") and part.endswith("]"):
            run = para.add_run(part)
            run.bold = True
            run.font.color.rgb = _COLOR_ORANGE
        else:
            para.add_run(part)


def _set_font(run, size_pt: int = 13, bold: bool = False,
              color: Optional[RGBColor] = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # Ensure East Asian font is also set (Windows compatibility)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _init_doc() -> Document:
    """Tạo Document mới với cài đặt lề và font mặc định chuẩn pháp lý VN."""
    doc = Document()

    # Lề trang chuẩn
    for section in doc.sections:
        section.left_margin  = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Style Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    # Style Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Style Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = _COLOR_BLUE

    return doc


def _add_line(doc: Document, line: str) -> None:
    """Phân tích một dòng và thêm paragraph tương ứng vào doc."""
    stripped = line.rstrip()

    # Dòng trống
    if not stripped:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        return

    # Markdown H2/H3
    m = _H3_RE.match(stripped)
    if m:
        p = doc.add_heading(m.group(1), level=2)
        return

    m = _H2_RE.match(stripped)
    if m:
        p = doc.add_heading(m.group(1), level=1)
        return

    # Section header ─── TEXT ───
    m = _SECTION_RE.match(stripped)
    if m:
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        _set_font(run, bold=True, color=_COLOR_BLUE)
        return

    # Dòng phân cách ─────────
    if _SEPARATOR_RE.match(stripped):
        p = doc.add_paragraph("─" * 48)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = _COLOR_GRAY
            run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        return

    # Bảng markdown | col | col |
    if _TABLE_ROW_RE.match(stripped):
        if _TABLE_SEP_RE.match(stripped):
            return   # bỏ dòng ---|---
        cells = [c.strip() for c in stripped.split("|")[1:-1]]
        table = doc.add_table(rows=1, cols=len(cells))
        table.style = "Table Grid"
        row = table.rows[0]
        for i, cell_text in enumerate(cells):
            cell = row.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            _add_formatted_runs(para, cell_text)
            for run in para.runs:
                _set_font(run)
        return

    # Bullet / danh sách
    m = _BULLET_RE.match(stripped)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        _add_formatted_runs(p, m.group(1))
        for run in p.runs:
            _set_font(run)
        return

    # Dòng quốc hiệu / tiêu đề căn giữa
    if _is_center_line(stripped):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped)
        _set_font(run)
        return

    # Dòng tiêu đề văn bản (ALL CAPS + từ khoá)
    if _is_title_line(stripped):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped)
        _set_font(run, size_pt=14, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        return

    # Đoạn văn thông thường (với **bold** và [placeholder])
    p = doc.add_paragraph()
    _add_formatted_runs(p, stripped)
    for run in p.runs:
        _set_font(run)


# ── Public API ─────────────────────────────────────────────────────────────────

def _doc_type_slug(doc_type: str) -> str:
    """Chuyển tên loại văn bản → slug ASCII cho tên file."""
    _MAP = {
        "biên bản vi phạm": "bien_ban_vi_pham",
        "đơn ly hôn":        "don_ly_hon",
        "hợp đồng lao động": "hop_dong_lao_dong",
        "hợp đồng thuê nhà": "hop_dong_thue_nha",
        "hợp đồng mua bán":  "hop_dong_mua_ban",
        "di chúc":           "di_chuc",
        "đơn khiếu nại":     "don_khieu_nai",
        "đơn tố cáo":        "don_to_cao",
        "công văn":          "cong_van",
        "biên bản":          "bien_ban",
        "quyết định":        "quyet_dinh",
    }
    lower = doc_type.lower().strip()
    for key, slug in _MAP.items():
        if key in lower:
            return slug
    # Fallback: bỏ dấu, giữ chữ cái/số
    nfkd = unicodedata.normalize("NFD", lower)
    ascii_s = "".join(c for c in nfkd if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]+", "_", ascii_s).strip("_") or "van_ban"


def text_to_docx(text: str, output_path: Path, doc_type: str = "") -> None:
    """Chuyển text/markdown → file .docx tại output_path."""
    doc = _init_doc()

    # Tách phần văn bản chính và phần CĂN CỨ PHÁP LÝ
    main_text = text
    legal_basis = ""
    split_markers = ["⚖️ CĂN CỨ PHÁP LÝ", "CĂN CỨ PHÁP LÝ ÁP DỤNG"]
    for marker in split_markers:
        idx = text.find(marker)
        if idx != -1:
            main_text   = text[:idx].rstrip()
            legal_basis = text[idx:]
            break

    # Thêm từng dòng của phần chính
    for line in main_text.split("\n"):
        _add_line(doc, line)

    # Thêm phần căn cứ pháp lý (nếu có) với page break nhỏ
    if legal_basis:
        doc.add_paragraph()
        for line in legal_basis.split("\n"):
            _add_line(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def export_draft(text: str, doc_type: str, export_dir: Path) -> Path:
    """Tạo file DOCX từ text soạn thảo. Trả về đường dẫn file đã lưu."""
    slug      = _doc_type_slug(doc_type)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"{slug}_{timestamp}.docx"
    out_path  = export_dir / filename
    text_to_docx(text, out_path, doc_type=doc_type)
    return out_path
